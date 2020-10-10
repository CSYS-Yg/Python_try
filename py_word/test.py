# 读取docx中的文本代码示例

import docx
import os, re

# 获取文档对象
# file = docx.Document("06丨 轻度痘痘 _ 外用药的挑选及使用.docx")
# print("段落数:" + str(len(file.paragraphs)))

# 输出段落编号及段落内容
# for i in range(len(file.paragraphs)):
#     print("段落数:" + str(i) + str(dir(file.paragraphs[i])))
# x = 140
# for i in range(x, x + 1):
#     p = file.paragraphs[i]
#     if len(p.text) > 0:
#         print(p.runs[0].font.color.rgb)
#         print(p.runs[0].font.bold)
#         print(p.runs[0].font.size.pt)


def test(filePath, result_path):
    for file in os.listdir(filePath):
        # 跳过非docx文件
        if ".docx" not in file:
            continue
        # 创建imgPath
        if not os.path.exists(result_path):
            os.makedirs(result_path)
        doc = docx.Document(filePath + file)  # 打开文件
        for rel in doc.part._rels:
            rel = doc.part._rels[rel]  # 获得资源
            if "image" not in rel.target_ref:
                continue
            imgName = re.findall("/(.*)", rel.target_ref)[0]
            with open(file + imgName, "wb") as f:
                f.write(rel.target_part.blob)


def get_pictures(word_path, result_path):
    """
    提取word文档内的图片
    :param word_path: word文件
    :param result_path: 结果目录
    :return:
    """
    doc = docx.Document(word_path)
    dict_rel = doc.part._rels
    for rel in dict_rel:
        rel = dict_rel[rel]
        print(rel.target_ref)
        if "image" in rel.target_ref:
            # 判断目录是否存在，创建目录
            if not os.path.exists(result_path):  # os.path.exists 检查路径是否存在
                os.makedirs(result_path)  # os.makedirs 创建路径
            img_name = re.findall("/(.*)", rel.target_ref)[0]
            print(img_name)
            word_name = os.path.splitext(word_path)[0]
            if os.sep in word_name:
                new_name = word_name.split('\\')[-1]
            else:
                new_name = word_name.split('/')[-1]
            img_name = f'{new_name}_{img_name}'
            with open(f'{result_path}/{img_name}', "wb") as f:
                f.write(rel.target_part.blob)


test("./", './')

# run 下枚举的属性
# ['__class__', '__delattr__', '__dict__', '__dir__', '__doc__', '__eq__','__format__', '__ge__', '__getattribute__', '__gt__', '__hash__',
#  '__init__', '__init_subclass__', '__le__', '__lt__', '__module__', '__ne__','__new__', '__reduce__', '__reduce_ex__', '__repr__',
#   '__setattr__', '__sizeof__', '__str__', '__subclasshook__', '__weakref__','_element', '_parent', '_r',
#   'add_break', 'add_picture', 'add_tab', 'add_text', 'bold', 'clear','element', 'font', 'italic', 'part', 'style', 'text', 'underline']

# from pydocx import PyDocX

# html = PyDocX.to_html("test.docx")

# f = open("test.html", 'w', encoding="utf-8")

# f.write(html)

# f.close()
